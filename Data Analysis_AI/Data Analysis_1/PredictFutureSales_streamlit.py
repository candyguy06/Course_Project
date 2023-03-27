# %% streamlit
# cd C:/Users/harry/Desktop/1091人工智慧實務/pj1/
# streamlit run PredictFutureSales_test1.py
import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
from docx import Document     #-- 小心: 必須 pip install python-docx 否則會錯

#%%
    

#%%

st.title("Predict Future Sales")
st.write("wei_ting-lin 2021/1/5")
#
# 資料來源:https://www.kaggle.com/c/competitive-data-science-predict-future-sales/data
# streamlit教學:https://blog.csdn.net/weixin_40575651/article/details/107603207

st.header("(KDD0)數據來源")
st.markdown("### 簡介")
st.write("資料來源: Kaggle-Predict Future Sales")
st.write("以2013.01~2015.10銷售數據預測2015.11每家商店賣出的產品數量")
st.write("檔案說明:")
st.markdown("""
            - sales_train.csv-2013年1月至2015年10月的每日曆史數據
            - test.csv-需要預測這些商店和產品在2015年11月的銷售額
            - items.csv-項目/產品補充訊息
            - item_categories.csv- 項目/類別補充訊息
            - shop.csv-店家補充訊息
            """)
#            - sample_submission.csv-正確格式的樣本
 


st.header("(KDD1)數據擷取")
st.markdown("### 讀入資料") 
import pandas as pd
path = "C:/Users/ben/Desktop/CYCU_AI/pj1/"
sales = pd.read_csv(path+'sales_train.csv')
items=pd.read_csv(path+'items.csv')
shops=pd.read_csv(path+'shops.csv')
item_cats=pd.read_csv(path+'item_categories.csv')
test=pd.read_csv(path+'test.csv')

#sales.head(3)
st.write("訓練數據:sales_train.csv(sales):")
st.write("sales.shape=",sales.shape) 
st.write("約290萬筆交易數據")
st.write("sales.head(5)=",sales.head(5)) 
st.markdown("""
            - date:商品售出日期
            - date_block_num:連續月份(2013年1月為0)
            - shop_id:此商品售出的店家
            - item_id:此商品ID
            - item_price:此商品標價
            - item_cnt_day-當日銷售的產品數量
            """)

st.write("品項數據:items.csv(items):")
st.write("items.shape=",items.shape) 
st.write("共有22170種品項")
st.write("items.head(5)=",items.head(5)) 
st.markdown("""
            - item_name:品項名稱
            - item_id:品項編碼(流水碼)
            - item_category_id:品項對應類別
            """)

st.write("類別數據:item_categories.csv(item_cats):")
st.write("item_cats.shape=",item_cats.shape) 
st.write("共有84種類別")
st.write("item_cats.head(5)=",item_cats.head(5)) 
st.markdown("""
            - item_category_name:類別名稱
            - item_category_id:類別編碼(流水碼)
            """)


st.write("通路數據:shops.csv(shops):")
st.write("shops.shape=",shops.shape) 
st.write("共有60家通路")
st.write("shops.head(5)=",shops.head(5)) 
st.markdown("""
            - shops_name:通路名稱
            - shops_id:通路編碼(流水碼)
            """)

st.write("預測數據:test.csv(test):")
st.write("test.shape=",shops.shape) 
st.write("共有214200個預測目標")
st.write("test.head(5)=",test.head(5)) 
st.markdown("""
            - ID:流水碼
            - shops_id:通路編碼
            - item_id:品項編碼
            """)


st.header("(KDD2)數據探索")
st.markdown("### (2-1)創建年/月欄位") 
st.write("因為要預測的銷售數量是以月為單位，所以獨立出來以利分析")
from datetime import datetime
sales['year'] = pd.to_datetime(sales['date']).dt.strftime('%Y')
sales['month'] = sales.date.apply(lambda x: datetime.strptime(x,'%d.%m.%Y').strftime('%m'))
sales.head(5)
st.write("sales.head(5)=",sales.head(5))


st.markdown("### (2-2)數據分布") 

import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns  #https://seaborn.pydata.org/tutorial/relational.html

st.set_option('deprecation.showPyplotGlobalUse', False) # PyplotGlobalUseWarning
ymgrouped = pd.DataFrame(sales.groupby(['year','month'])['item_cnt_day'].sum().reset_index())
st.write("歷年銷售數據(銷售量/月):")
sns.pointplot(x='month', y='item_cnt_day', hue='year', data=ymgrouped)
st.pyplot()
st.write("以整體銷售量數據來看")
st.write("1.售出的商品數量有逐年減少的趨勢")
st.write("2.在接近年尾時商品賣得較好")
st.write("2.年尾是銷售量最高")


grouped_price = pd.DataFrame(sales.groupby(['year','month'])['item_price'].mean().reset_index())
sns.pointplot(x='month', y='item_price', hue='year', data=grouped_price)
st.write("歷年銷售數據(銷售金額/月):")
st.pyplot()
st.write("以整體銷售額數據來看")
st.write("1.售出的商品價格逐年提高")
st.write("2.每年8-9月金額增加幅度大")
st.write("3.年尾是銷售額最高")

ts=sales.groupby(["date_block_num"])["item_cnt_day"].sum()
ts.astype('float')
plt.figure(figsize=(16,8))
plt.title('Total Sales of the whole time period')
plt.xlabel('Time')
plt.ylabel('Sales')
plt.plot(ts);
st.write("連續月份銷量數據(銷售量/連續月份):")
st.pyplot()
st.write("以連續月份銷售量來看")
st.write("1.銷售量逐年下降")
st.write("2.春夏之間銷量低迷")
st.write("3.年底是銷量高峰")


sns.jointplot(x="item_cnt_day", y="item_price", data=sales, height=8)
st.write("價格與銷量分布圖")
st.pyplot()
st.write("以價格與銷量分布圖來看")
st.write("1.價格帶主要落在0~50000之間")
st.write("2.單筆交易數量主要落在0~500之間")
st.write("3.價格和單筆件數離異點")



st.markdown("### (2-3)數據清理")

st.write("數據清理前資料大小:",sales.shape)
sales = sales.query('item_price > 0')
test = pd.read_csv(path+'test.csv')
# 測試數據內沒有的店家及商品不進行分析
sales = sales[sales['shop_id'].isin(test['shop_id'].unique())]
sales = sales[sales['item_id'].isin(test['item_id'].unique())]
sales = sales.query('item_cnt_day >= 0 and item_cnt_day <= 125 and item_price < 75000')
st.write("數據清理後資料大小:",sales.shape)
st.write("清除條件 : ")
st.write("1.單價小於0(退貨)")
st.write("2.預測目標不包含的通路及品項")
st.write("3.單筆件數大於125以及單筆金額大於75000(離異點)")

sns.jointplot(x="item_cnt_day", y="item_price", data=sales, height=8)
st.write("清洗後的銷量價格分布:")
st.pyplot()
st.write("1.價格主要落在0~10000之間")
st.write("2.單筆交易數量主要落在0~20之間")
st.write("3.單價越低數量越高")
st.write("4.單價20000附近的價格帶也有高銷量")


ymcleaned = pd.DataFrame(sales.groupby(['year','month'])['item_cnt_day'].sum().reset_index())
sns.pointplot(x='month', y='item_cnt_day', hue='year', data=ymcleaned)
st.write("資料清理後歷年銷售數據(數量/月):")
st.pyplot()
st.write("以要預測的商品來看")
st.write("1.是銷售數量逐年增加的產品")
st.write("2.同樣的在接近年尾時商品賣得較好")

st.markdown("### 小節") 
st.write("1.年尾銷售量銷售額都比較高")
st.write("2.銷售量逐年下降")
st.write("3.銷售額逐年上升")
st.write("4.銷售量下降,銷售額卻上升 -> 價格提高 -> 俄羅斯金融風暴")

st.header("(KDD3)數據轉換")
st.markdown("### 將sales數據依通路以及品項列出連續33個月的銷售量")
monthly_sales=sales.groupby(["date_block_num","shop_id","item_id"])[
    "date_block_num","date","item_price","item_cnt_day"].agg({"date_block_num":'mean',"date":["min",'max'],"item_price":"mean","item_cnt_day":"sum"})
sales_data_flat = monthly_sales.item_cnt_day.apply(list).reset_index()
sales_data_flat = pd.merge(test,sales_data_flat,on = ['item_id','shop_id'],how = 'left')
sales_data_flat.fillna(0,inplace = True)
sales_data_flat.drop(['shop_id','item_id'],inplace = True, axis = 1)
sales_data_flat.head(5)
#   ID  date_block_num  sum
#0   0            20.0  1.0
#1   0            22.0  1.0
#2   0            23.0  2.0
#3   0            24.0  2.0
#4   0            28.0  1.0
pivoted_sales = sales_data_flat.pivot_table(index='ID', columns='date_block_num',fill_value = 0,aggfunc='sum' )
st.write("pivoted_sales.head(5)=",pivoted_sales.head(5)) 

st.write("將數據轉換成以通路品項為行,33個連續月份銷量為列的資料格式")


st.markdown("### 通路聚類")
sales2 = pd.read_csv(path+'sales_train.csv')
Sv1 = sales2.groupby(['shop_id'])['item_cnt_day'].sum().reset_index()
Sv2 = sales2.groupby(['shop_id'])['item_price'].sum().reset_index()
Sv1['item_price'] = Sv2['item_price']

from sklearn.cluster import KMeans
#from sklearn_extra.cluster import KMedoids
KM=KMeans(n_clusters=3,init='random',random_state=5)
KM.fit(Sv1)
k = KM.predict(Sv1)
from pandas import DataFrame
df = DataFrame (k,columns=['group'])
Sv1['group'] = df['group']
a = Sv1[Sv1['group'] == 0 ]
b = Sv1[Sv1['group'] == 1 ] 
c = Sv1[Sv1['group'] == 2 ]
a['one'] = a['item_price'] / a['item_cnt_day']
b['one'] = b['item_price'] / b['item_cnt_day']
c['one'] = c['item_price'] / c['item_cnt_day']
st.write("高單價通路:",a.head(len(a['shop_id'])))
st.write("通路數量: {} ".format(len(a["one"])), style='List Bullet' )
st.write("單價區間: {} - {}".format(min(a["one"]), max(a["one"])), style='List Bullet' )
st.write("單價平均: {:.2f} ".format((a["one"].mean())), style='List Bullet' )
st.write("銷量區間: {} - {}".format(min(a["item_cnt_day"]), max(a["item_cnt_day"])), style='List Bullet' )
st.write("銷量平均: {:.2f} ".format((a["item_cnt_day"].mean())), style='List Bullet' )
st.write("總銷售額區間: {} - {}".format(min(a["item_price"]), max(a["item_price"])), style='List Bullet' )
st.write("總銷售額平均: {:.2f} ".format((a["item_price"].mean())), style='List Bullet' )


st.write("高銷量通路:",b.head(len(b['shop_id'])))
st.write("通路數量: {} ".format(len(b["one"])), style='List Bullet' )
st.write("單價區間: {} - {}".format(min(b["one"]), max(b["one"])), style='List Bullet' )
st.write("單價平均: {:.2f} ".format((b["one"].mean())), style='List Bullet' )
st.write("銷量區間: {} - {}".format(min(b["item_cnt_day"]), max(b["item_cnt_day"])), style='List Bullet' )
st.write("銷量平均: {:.2f} ".format((b["item_cnt_day"].mean())), style='List Bullet' )
st.write("總銷售額區間: {} - {}".format(min(b["item_price"]), max(b["item_price"])), style='List Bullet' )
st.write("總銷售額平均: {:.2f} ".format((b["item_price"].mean())), style='List Bullet' )


st.write("一般通路:",c.head(len(c['shop_id'])))
st.write("通路數量: {} ".format(len(c["one"])), style='List Bullet' )
st.write("單價區間: {} - {}".format(min(c["one"]), max(c["one"])), style='List Bullet' )
st.write("單價平均: {:.2f} ".format((c["one"].mean())), style='List Bullet' )
st.write("銷量區間: {} - {}".format(min(c["item_cnt_day"]), max(c["item_cnt_day"])), style='List Bullet' )
st.write("銷量平均: {:.2f} ".format((c["item_cnt_day"].mean())), style='List Bullet' )
st.write("總銷售額區間: {} - {}".format(min(c["item_price"]), max(c["item_price"])), style='List Bullet' )
st.write("總銷售額平均: {:.2f} ".format((c["item_price"].mean())), style='List Bullet' )

st.write("以通路聚類來看")
st.write("1.高銷量通路平均總銷售額較高")
st.write("2.高單價通路雖然單價比一般通路高，但平均銷售數也比較一般通路高")
st.write("3.高銷量通路平均總銷售額明顯較高")



st.header("(KDD4)模型訓練")
st.markdown("### (4-1)訓練集與測試集")
X_train = np.expand_dims(pivoted_sales.values[:,:-1],axis = 2)
y_train = pivoted_sales.values[:,-1:]
X_test = np.expand_dims(pivoted_sales.values[:,1:],axis = 2)
st.write("X_train:")
st.write("X_train.shape=",X_train.shape) 

st.write("y_train:")
st.write("y_train.shape=",y_train.shape) 

st.write("X_test:")
st.write("X_test.shape=",X_test.shape) 



st.markdown("### (4-2)模型建立")
from keras.models import Sequential
from keras.layers import LSTM,Dense,Dropout
from keras.models import load_model, Model
import matplotlib.image as mpimg
from keras import layers
from tensorflow import keras
st.write("LSTM模型") 
#sales_model = Sequential()
#sales_model.add(LSTM(units = 64,input_shape = (33,1)))
##sales_model.add(LSTM(units = 64,activation='relu'))
#sales_model.add(Dropout(0.5))
#sales_model.add(Dense(1))
#
#sales_model.compile(loss = 'mse',optimizer = 'adam', metrics = ['mean_squared_error'])
#st.write(sales_model.summary())
sales_model = keras.models.load_model("my_model")
img1=mpimg.imread('m1.jpg')
img2=mpimg.imread('mt1.jpg')
imgplot1 = plt.imshow(img1)
imgplot2 = plt.imshow(img2)
st.image("m1.jpg", use_column_width=True)
st.image("mt1.jpg", use_column_width=True)            

st.write("回歸模型") 
#sales_model2 = Sequential()
##sales_model2.add(LSTM(units = 64,input_shape = (33,1)))
#sales_model2.add(layers.Dense(units = 64, activation = 'relu',input_shape = (X_train.shape[1],1)))
##sales_model.add(LSTM(units = 64,activation='relu'))
#sales_model2.add(layers.Dense(64,activation = 'relu'))
##sales_model2.add(Dropout(0.5))
#sales_model2.add(Dense(1))
#sales_model2.compile(loss = 'mse',optimizer = 'rmsprop', metrics = ['mean_squared_error'])
##st.write(sales_model2.summary())
sales_model2 = keras.models.load_model("my_model2")
img3=mpimg.imread('m2.jpg')
img4=mpimg.imread('mt2.jpg')
imgplot3 = plt.imshow(img3)
imgplot4 = plt.imshow(img4)
st.image("m2.jpg", use_column_width=True)
st.image("mt2.jpg", use_column_width=True)            



st.header("(KDD5)預測")
st.markdown("### 預測結果") 
    
#sales_model.fit(X_train,y_train,batch_size = 4096,epochs = 10)
submission_output = sales_model.predict(X_test)
submission = pd.DataFrame({'item_cnt_month':submission_output.ravel()})
#submission.to_csv('submission_stacked.csv',index = False)
st.write("submission1=",submission.head(5)) 

#df1 = pd.DataFrame()
#df1['item_cnt_day'] = submission['item_cnt_month']
#df1['date'] = "01.11.2015"
#df1['year'] = pd.to_datetime(df1['date']).dt.strftime('%Y')
#df1['month'] = df1.date.apply(lambda x: datetime.strptime(x,'%d.%m.%Y').strftime('%m')) #another way for same thing
#sale2 = pd.DataFrame()
#sale2['date'] = sales['date']
#sale2['year'] = sales['year']
#sale2['month'] = sales['month']
#sale2['item_cnt_day'] = sales['item_cnt_day']
#sale2 = sale2.append(df1)
#ymcleaned2 = pd.DataFrame(sale2.groupby(['year','month'])['item_cnt_day'].sum().reset_index())
#sns.pointplot(x='month', y='item_cnt_day', hue='year', data=ymcleaned2)
#st.pyplot()
st.write("lstm預測後歷年銷售數據(數量/月):")
from PIL import Image
image = Image.open('predict.jpg')
st.image(image, caption='Sunrise by the mountains',use_column_width=True)
st.write("預測11月銷量些微下降")
st.write("依歷年數據來看即便沒有大幅度提升也應該要有上升趨勢")



##sales_model2.fit(X_train,y_train,batch_size = 4096,epochs = 10)
#submission_output = sales_model2.predict(X_test)
#submission = pd.DataFrame({'item_cnt_month':submission_output.ravel()})
##submission.to_csv('submission_stacked.csv',index = False)
#st.write("submission2=",submission.head(5))

 



def doc_table(dx,TT):  # called by gen???RPT?() to generate table in .docx
    # Jia-Sheng Heh, 06/25/2020
    DT = dx.add_table(rows=1, cols=TT.shape[1]+1, style='Light Grid Accent 1' )
    Hcell = DT.rows[0].cells;    Hcell[0].text = "index"
    for k in np.arange(len(TT.columns)):    Hcell[k+1].text = TT.columns[k]
    for i in list(TT.index):
        Rcell = DT.add_row().cells;     Rcell[0].text = str(i)
        SS = [str(x) for x in TT.loc[i]]
        for j in np.arange(len(TT.columns)):    Rcell[j+1].text = SS[j]
def genOpAnalysis01(X,rptFile,a,b,c):   # to generate Operational-Report-01.docx (customer values/retention/return)
    # Jia-Sheng Heh, 06/28/2020 temporarily
    import numpy as np
    import pandas as pd
    import plotly.express as px
    # from plotly.offline import plot  
    from docx import Document     #-- 小心: 必須 pip install python-docx 否則會錯
    from docx import shared
    
    dx = Document()
    dx.add_heading('銷售探索', 0)  #----- REPORT INITIALIZATION
    dx.add_paragraph('--這是由 RDS系統自動生成的報告')
#    dx.add_paragraph(" ".join(['--日期: ',datetime.datetime.now().strftime("%m/%d/%Y")]) )
    dx.add_paragraph('--作者: lin wai ting')
    dx.add_heading('1.(KDD1) 擷取數據 (-->X)', level=1)  #----- Section 1
    dx.add_paragraph("數據日期: {} - {}".format(min(X["date"]), max(X["date"])), style='List Bullet' )
    dx.add_paragraph("交易筆數: {} 筆紀錄".format(X.shape[0]), style='List Bullet' )
    dx.add_paragraph("交易數據樣貌 --", style='List Bullet' )
    doc_table(dx, X.head(3))
    dx.add_heading('2.(KDD2) 數據探索 (X)', level=1)  #----- Section 2

    dx.add_paragraph("歷年銷售數據(銷售量/月): --", style='List Bullet' )
    dx.add_picture('mcnt.jpg',width=shared.Cm(15) )    
    dx.add_paragraph("以整體銷售量數據來看")
    dx.add_paragraph("1.售出的商品數量有逐年減少的趨勢")
    dx.add_paragraph("2.在接近年尾時商品賣得較好")
    dx.add_paragraph("3.年尾是銷售量最高")

    dx.add_paragraph("歷年銷售數據(銷售金額/月): --", style='List Bullet' )
    dx.add_picture('mpri.jpg',width=shared.Cm(15) )    
    dx.add_paragraph("以整體銷售額數據來看")
    dx.add_paragraph("1.售出的商品價格逐年提高")
    dx.add_paragraph("2.每年8-9月金額增加幅度大")
    dx.add_paragraph("3.年尾是銷售額最高")
    
    dx.add_paragraph("連續月份銷量數據(銷售量/連續月份): --", style='List Bullet' )
    dx.add_picture('msale.jpg',width=shared.Cm(15) )    
    dx.add_paragraph("以連續月份銷售量來看")
    dx.add_paragraph("1.銷售量逐年下降")
    dx.add_paragraph("2.春夏之間銷量低迷")
    dx.add_paragraph("3.年底是銷量高峰")

    dx.add_paragraph("價格與銷量分布圖: --", style='List Bullet' )
    dx.add_picture('cp.jpg',width=shared.Cm(15) )    
    dx.add_paragraph("以價格與銷量分布圖來看")
    dx.add_paragraph("1.價格帶主要落在0~50000之間")
    dx.add_paragraph("2.單筆交易數量主要落在0~500之間")


    dx.add_paragraph("小結")    
    dx.add_paragraph("1.年尾銷售量銷售額都比較高")
    dx.add_paragraph("2.銷售量逐年下降")
    dx.add_paragraph("3.銷售額逐年上升")    
    dx.add_paragraph("4.銷售量下降,銷售額卻上升 -> 價格提高 -> 俄羅斯金融風暴")    
    

    dx.add_heading('2.(KDD3) 通路聚類 (X)', level=1)  #----- Section 3
    dx.add_paragraph("通路A: 共 {} 家通路".format(len(a["one"])), style='List Bullet' )
    doc_table(dx, a.head(len(a['shop_id'])))
    dx.add_paragraph("高單價通路:")
    dx.add_paragraph("單價區間: {} - {}".format(min(a["one"]), max(a["one"])), style='List Bullet' )
    dx.add_paragraph("單價平均: {:.2f} ".format((a["one"].mean())), style='List Bullet' )
    dx.add_paragraph("銷量區間: {} - {}".format(min(a["item_cnt_day"]), max(a["item_cnt_day"])), style='List Bullet' )
    dx.add_paragraph("銷量平均: {:.2f} ".format((a["item_cnt_day"].mean())), style='List Bullet' )
    dx.add_paragraph("總銷售額區間: {} - {}".format(min(a["item_price"]), max(a["item_price"])), style='List Bullet' )
    dx.add_paragraph("總銷售額平均: {:.2f} ".format((a["item_price"].mean())), style='List Bullet' )
    
    dx.add_paragraph("通路B: 共 {} 家通路".format(len(b["one"])), style='List Bullet' )
    doc_table(dx, b.head(len(b['shop_id'])))
    dx.add_paragraph("高銷量通路:")
    dx.add_paragraph("單價區間: {} - {}".format(min(b["one"]), max(b["one"])), style='List Bullet' )
    dx.add_paragraph("單價平均: {:.2f} ".format((b["one"].mean())), style='List Bullet' )
    dx.add_paragraph("銷量區間: {} - {}".format(min(b["item_cnt_day"]), max(b["item_cnt_day"])), style='List Bullet' )
    dx.add_paragraph("銷量平均: {:.2f} ".format((b["item_cnt_day"].mean())), style='List Bullet' )
    dx.add_paragraph("總銷售額區間: {} - {}".format(min(b["item_price"]), max(b["item_price"])), style='List Bullet' )
    dx.add_paragraph("總銷售額平均: {:.2f} ".format((b["item_price"].mean())), style='List Bullet' )

    dx.add_paragraph("通路C: 共 {} 家通路".format(len(c["one"])), style='List Bullet' )
    doc_table(dx, c.head(len(c['shop_id'])))
    dx.add_paragraph("一般通路:")
    dx.add_paragraph("單價區間: {} - {}".format(min(c["one"]), max(c["one"])), style='List Bullet' )
    dx.add_paragraph("單價平均: {:.2f} ".format((c["one"].mean())), style='List Bullet' )
    dx.add_paragraph("銷量區間: {} - {}".format(min(c["item_cnt_day"]), max(c["item_cnt_day"])), style='List Bullet' )
    dx.add_paragraph("銷量平均: {:.2f} ".format((c["item_cnt_day"].mean())), style='List Bullet' )
    dx.add_paragraph("總銷售額區間: {} - {}".format(min(c["item_price"]), max(c["item_price"])), style='List Bullet' )
    dx.add_paragraph("總銷售額平均: {:.2f} ".format((c["item_price"].mean())), style='List Bullet' )

    dx.add_paragraph("以通路聚類來看")
    dx.add_paragraph("1.高銷量通路平均總銷售額較高")
    dx.add_paragraph("2.高單價通路雖然單價比一般通路高，但平均銷售數也比較一般通路高")
    dx.add_paragraph("3.高銷量通路平均總銷售額明顯較高")


    dx.save(rptFile)   #----- REPORT GENREATION



#if st.sidebar.checkbox("word報告 (sales-->sales.docx) --"):
st.header("生成word檔")
genOpAnalysis01(sales,"sales.docx",a,b,c)   
st.write("  ** 己產生報告(sales.docx)")



# %%
